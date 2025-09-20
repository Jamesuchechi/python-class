from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import os

# === FILES ===
input_file = "/Users/Dev Harry/Documents/Techminds_Academy_Marketers_Enhanced.docx"
logo_file = "/Users/Dev Harry/Pictures/logo.webp"
output_file = "/Users/Dev Harry/Documents/Techminds_Academy_With_Images.docx"

# converting all images to png
def ensure_png(image_path):
    """Convert any image to PNG and return the safe path."""
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"❌ File not found: {image_path}")
    base, _ = os.path.splitext(image_path)
    new_path = base + ".png"
    Image.open(image_path).convert("RGB").save(new_path, "PNG")
    return new_path

# Convert logo to PNG
logo_png = ensure_png(logo_file)

# Map program sections to icons 
icon_map = {
    "Web Design/Development": "/Users/Dev Harry/Pictures/software.jpg",
    "Data Science": "/Users/Dev Harry/Pictures/icon_datascience.png",
    "Digital Marketing": "/Users/Dev Harry/Pictures/graphic.jpg",
    "Graphic Design": "/Users/Dev Harry/Pictures/admission-picture.jpg",
    "Cybersecurity": "/Users/Dev Harry/Pictures/web.png",   
}

# Convert all icons to PNG
safe_icon_map = {prog: ensure_png(path) for prog, path in icon_map.items()}

# Load the document
doc = Document(input_file)

# Insert logo at the very top, centered
first_para = doc.paragraphs[0].insert_paragraph_before()
first_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
first_para.add_run().add_picture(logo_png, width=Inches(2))

# Insert program icons before their section headings
for para in list(doc.paragraphs):
    text = para.text.strip()
    for program, icon_file in safe_icon_map.items():
        if program in text and os.path.exists(icon_file):
            run = para.insert_paragraph_before().add_run()
            run.add_picture(icon_file, width=Inches(1.2))
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            break

# Save the new document
doc.save(output_file)
print(f"✅ Done! New document saved as {output_file}")
